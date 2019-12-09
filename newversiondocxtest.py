from docx import Document
from random import randint

w = Document('pss.docx')
mas=[]
tables_size =len( w.tables)
rows=len(w.tables[1].rows)
k1={}
k2={}
k3={}
z=None
i=0
d1=[x for x in range(1,rows,7)]
d2=[x for x in range(1,len(w.tables[2].rows),8)]
d3=[x for x in range(1,len(w.tables[3].rows),9)]
col=0
while True:
  i+=1
  while True:
    z=randint(1,len(w.tables[i].rows))
    if i==1:
        for j in d1:
            if j>z:
                z=j
                col+=1
                k1[w.tables[i].rows[z].cells[-1].text]=[w.tables[i].rows[z+1].cells[-1].text,set(w.tables[i].rows[x].cells[-1].text for x in range(z+2,z+6) )]
                break
    if i==2:
        for j in d2:
            if j>z:
                z=j
                col+=1
                k2[w.tables[i].rows[z].cells[-1].text]=[w.tables[i].rows[z+1].cells[-1].text,w.tables[i].rows[z+2].cells[-1].text,set(w.tables[i].rows[x].cells[-1].text for x in range(z+3,z+7) )]
                break
    if i==3:
        for j in d3:
            if j>z:
                z=j
                col+=1
                k3[w.tables[i].rows[z].cells[-1].text]=[w.tables[i].rows[z+1].cells[-1].text,w.tables[i].rows[z+2].cells[-1].text,w.tables[i].rows[z+3].cells[-1].text,set(w.tables[i].rows[x].cells[-1].text for x in range(z+4,z+8) )]
                break
    if col==10:
        i+=1
    elif col==20:
        i+=1
    elif col==30:
        break
  break
s=0
f=0
k=None
b=-1
kq=None
vb=-1
for i in k1.items():
    print(i[0])
    f=randint(1,4)
    s=0
    kq=None
    k=None
    b=-1
    vb=-1
    for j in i:
        s+=1
        if isinstance(j,list):
            while True:
                if s==f:
                    k=i[1][1]
                    print(f'Ответ номер {s} : {k}')
                else:
                    b+=1
                    for h in j[1]:
                         vb+=1
                         if vb==b:
                             print(f'Ответ номер {s} : {h}')
                if len(j)==b:
                    c=int(input('Введите ответ : '))
                    if c==f:
                        print('yes')
                    else:
                       print('no')
                    break
                s+=1
                vb=-1   
            
        
            
        
    
            
        

    
    
            
    
    


            
        
        
    
    
        
        
    
    
    


    
    

        
            
    
    



