from docx import Document
from random import randint

wordDoc = Document('pss.docx')
mas=[]
tables_size =len( wordDoc.tables)
"""
for i in range(tables_size):
    mas.append(len(wordDoc.tables[i].rows))
    for j in range(mas[i]):
        print(wordDoc.tables[i].rows[j].cells[-1].text)
print(mas)
"""
rows=len(wordDoc.tables[1].rows)
"""
print(wordDoc.tables[1].rows[0].cells[-1].text)
"""
k1=[]
c1=5
k2=[]
k3=[]
f=None
z=0
d1=[x for x in range(1,rows,7)]
d2=[x for x in range(1,len(wordDoc.tables[2].rows),8)]
d3=[x for x in range(1,len(wordDoc.tables[3].rows),9)]
otv1=[]
print(d1)
count=0
sa=0
"""
print(wordDoc.tables[3].rows[28].cells[-1].text)

"""



while(True):
    f=randint(1,tables_size)
    if f==1:
        z=randint(1,rows)
        for g in d1:
            if g>z:
                z=g
                break
        k1.append(wordDoc.tables[f].rows[z].cells[-1].text)
        for i in range(z+1,z+5):
            otv1.append(wordDoc.tables[f].rows[i].cells[-1].text)
        count+=1
    elif f==2:
        z=randint(1,len(wordDoc.tables[2].rows))
        for g in d2:
            if g>z:
                z=g
                break
        k2.append(wordDoc.tables[f].rows[z].cells[-1].text)
            
        count+=1
    elif f==3:
        z=randint(1,len(wordDoc.tables[2].rows))
        for g in d3:
            if g>z:
                z=g
                break
                
        k3.append(wordDoc.tables[f].rows[1].cells[-1].text)
        count+=1
    if count==30:
        break
print(otv1)
bv=0
qa=0
print('-----------------------')
"""
for i in range(len(k1)):
    print(f'вопрос {i}')
    print(f'сам вопрос звучит : {k1[i]}')
    for j in otv1[qa:qa+5]:
        bv+=1
        print(f'ответ номер {bv} : {j}')
    bv=0
    qa+=5
"""

"""
nb=0
col=0
v=[]
nax=0
fh=[]
for i in range(len(k1)):
    print(f'вопрос {i}')
    print(f'сам вопрос звучит : {k1[i]}')
    v=otv1[qa:qa+5]
    for j in range(qa,qa+5):
        while True:
           nb=randint(0,len(v))
           if nb not in fh:
              fh+=[nb]
              col+=1
              h=v.pop(nb)
              if nb==qa:
                nax=col
              break
           else :
                continue
        print(f' ответ номер {col} : {h}')
    qa+=5
    col=0
    fh=[]
    nax=0
"""
q=None
pr=None
m=[]
hj=True
bl=0
col=0
otv=None
q=0
for i,j in enumerate(k1):
    print(f' вопрос номер {i} :  {j}')
    m=[]
    otv=None
    hj=True
    col=0
    bl=0
    while True:
        pr=otv1[q]
        while True:
            z=randint(q,q+5)
            if z not in m:
                col+=1
                if z==q:
                    bl=col   
                m+=[z]
                print(f'ответ номер {col} : {otv1[z]}')
                if len(m)==5:
                    break
            else:
                continue
        print(bl)
        otvet=int(input('Введите цифру ответа : '))
        if otvet==bl:
            print('Правильно')
        else:
            print('Не правильно ')
        q+=5
        break

            
        
        
    
    
        
        
    
    
    


    
    

        
            
    
    



