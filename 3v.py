from docx import Document
from random import randint
k=set()
wordDoc = Document('us.docx')
global mas
mas=None
global h
h=0
d=[randint(151,200) for x in range(20)]
tables_size =wordDoc.tables
print(wordDoc.tables[149].rows[0].cells[0].text)
for i in d:
     mas=[]
     for j in range(1,8):
         if wordDoc.tables[i].rows[j].cells[0].text == '1':
                  k.add('f'+wordDoc.tables[i].rows[j].cells[-1].text)

             
         else:
             k.add(wordDoc.tables[i].rows[j].cells[1].text)
     print(wordDoc.tables[i].rows[0].cells[-1].text)
     for z in k:
         h+=1
         if z[0]=='f':
             print(f'ответ {h} : {z[1:len(z)]}')
             mas.append(h)     
         else:
                print(f'ответ {h} : {z}')
     n=[int(x) for x in input('Введите ответ : ').split()]
     if n[0] in mas  and n[1] in mas  and n[2] in mas :
         print('Все верно ')
     else:
        print('правильно так : ')
        print(mas)
        
     h=0
     k=set()
     mas=None            
            
        
    
        
         
        

